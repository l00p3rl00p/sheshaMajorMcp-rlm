"""Office document parser for .docx files."""

from pathlib import Path

from docx import Document

from shesha.models import ParsedDocument


class DocxParser:
    """Parser for Word .docx files using python-docx."""

    def can_parse(self, path: Path, mime_type: str | None = None) -> bool:
        """Check if this parser can handle the given file."""
        return path.suffix.lower() == ".docx"

    def parse(
        self,
        path: Path,
        include_line_numbers: bool = False,
        file_path: str | None = None,
    ) -> ParsedDocument:
        """Parse a .docx file and return a ParsedDocument.

        Args:
            path: Path to the file to parse.
            include_line_numbers: Ignored for DOCX files.
            file_path: Ignored for DOCX files.
        """
        doc = Document(str(path))
        parts: list[str] = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))
            if table_text:
                parts.append("\n".join(table_text))

        content = "\n\n".join(parts)

        return ParsedDocument(
            name=path.name,
            content=content,
            format="docx",
            metadata={"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)},
            char_count=len(content),
            parse_warnings=[],
        )
